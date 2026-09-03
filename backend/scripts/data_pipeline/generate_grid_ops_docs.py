"""Generate a realistic GridOps document corpus grounded in the seeded SQL data.

The 8 hand-authored documents in seed/docs/true_data/ are the *procedural /
conceptual* layer (how the utility operates). This script generates the
*operational records* layer on top of them — the per-asset profiles, outage
post-mortems, inspection reports and periodic reliability summaries that make
up the bulk of a real utility's document store.

Every document is generated from actual rows in the grid_ops tables, so a
question like "what caused the outage on feeder X" is answerable from BOTH the
documents (RAG) and the database (Text2SQL) — which is what makes the hybrid
route meaningful.

Output formats are deliberately mixed (.txt/.md/.html/.docx/.pdf) to exercise
the full Docling ingestion path.

Usage:
    uv run python scripts/data_pipeline/generate_grid_ops_docs.py [--force]

Writes: seed/docs/true_data/operational_records/
"""

from __future__ import annotations

import argparse
import random
import shutil
from datetime import datetime
from pathlib import Path

import psycopg2
import psycopg2.extras

from app.config import settings

SEED = 1337
OUT_DIR = (
    Path(__file__).parent.parent.parent / "seed" / "docs" / "true_data" / "operational_records"
)

# How many documents of each type to emit.
N_SUBSTATION_PROFILES = 50   # one per substation
N_OUTAGE_POSTMORTEMS = 150
N_TRANSFORMER_INSPECTIONS = 100
N_FEEDER_SUMMARIES = 80
N_REGIONAL_REPORTS = 30
N_STORM_REPORTS = 12

# Format mix — .pdf/.docx are slowest to parse, so keep them a minority.
FORMAT_WEIGHTS = {".txt": 45, ".md": 20, ".html": 20, ".docx": 10, ".pdf": 5}

MONTHS = [
    "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December",
]

INSPECTOR_TITLES = ["Substation Technician", "Field Engineer", "Asset Inspector", "Relay Technician"]
OIL_CONDITIONS = ["clear, within spec", "slight discoloration, within spec",
                  "darkened — schedule dielectric test", "clear, recently filtered"]
BUSHING_CONDITIONS = ["no visible tracking", "minor surface contamination",
                      "clean, gasket intact", "light corrosion at flange"]
WEATHER = ["clear", "light rain", "high wind advisory", "ice accumulation",
           "extreme heat advisory", "thunderstorms"]
STORM_NAMES = ["Winter Storm Halcyon", "Tropical Depression Meridian", "Derecho Event",
               "Ice Storm Brindle", "Nor'easter Calloway", "Heat Dome Event",
               "Windstorm Pallas", "Winter Storm Ovid", "Tropical Storm Ketch",
               "Severe Convective Outbreak", "Ice Storm Verity", "Windstorm Tarn"]


# ---------------------------------------------------------------------------
# DB access
# ---------------------------------------------------------------------------


def fetch_all(conn, sql: str, params: tuple = ()) -> list[dict]:
    with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
        cur.execute(sql, params)
        return [dict(r) for r in cur.fetchall()]


# ---------------------------------------------------------------------------
# Format writers
# ---------------------------------------------------------------------------


def write_txt(path: Path, title: str, blocks: list[str]) -> None:
    path.write_text(title + "\n\n" + "\n\n".join(blocks), encoding="utf-8")


def write_md(path: Path, title: str, blocks: list[str]) -> None:
    body = f"# {title}\n\n" + "\n\n".join(blocks)
    path.write_text(body, encoding="utf-8")


def write_html(path: Path, title: str, blocks: list[str]) -> None:
    paras = "\n".join(f"  <p>{b}</p>" for b in blocks)
    html = (
        "<!doctype html>\n<html><head><meta charset='utf-8'>"
        f"<title>{title}</title></head>\n<body>\n  <h1>{title}</h1>\n{paras}\n</body></html>\n"
    )
    path.write_text(html, encoding="utf-8")


def write_docx(path: Path, title: str, blocks: list[str]) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    for b in blocks:
        doc.add_paragraph(b)
    doc.save(str(path))


def write_pdf(path: Path, title: str, blocks: list[str]) -> None:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(path), pagesize=LETTER)
    flow = [Paragraph(title, styles["Title"]), Spacer(1, 12)]
    for b in blocks:
        flow.append(Paragraph(b, styles["BodyText"]))
        flow.append(Spacer(1, 8))
    doc.build(flow)


WRITERS = {
    ".txt": write_txt,
    ".md": write_md,
    ".html": write_html,
    ".docx": write_docx,
    ".pdf": write_pdf,
}


def emit(rng: random.Random, slug: str, title: str, blocks: list[str]) -> str:
    ext = rng.choices(list(FORMAT_WEIGHTS), weights=list(FORMAT_WEIGHTS.values()))[0]
    path = OUT_DIR / f"{slug}{ext}"
    WRITERS[ext](path, title, blocks)
    return path.name


# ---------------------------------------------------------------------------
# Document builders
# ---------------------------------------------------------------------------


def build_substation_profile(rng: random.Random, sub: dict, feeders: list[dict],
                             transformers: list[dict], outages: list[dict]) -> tuple[str, str, list[str]]:
    name = sub["name"]
    total_capacity = sum(float(t["capacity_mva"]) for t in transformers)
    energized = [f for f in feeders if f["status"] == "energized"]
    commissioned = sub["commissioned_at"].strftime("%B %Y")

    blocks = [
        f"Substation {name} is a {sub['substation_type']} substation located in the "
        f"{sub['region']} region, operating at a {sub['voltage_class_kv']} kV voltage class. "
        f"It was commissioned in {commissioned} and its current operating status is "
        f"{sub['status']}.",

        f"The substation serves {len(feeders)} distribution feeders, of which {len(energized)} "
        f"are currently energized. Installed transformer capacity across {len(transformers)} "
        f"units totals {total_capacity:.2f} MVA. "
        + (
            f"The largest single unit is {max(transformers, key=lambda t: float(t['capacity_mva']))['asset_tag']} "
            f"at {max(float(t['capacity_mva']) for t in transformers):.2f} MVA."
            if transformers else "No transformer assets are currently registered to this substation."
        ),
    ]

    if feeders:
        listing = "; ".join(
            f"{f['name']} ({f['voltage_kv']} kV, {f['capacity_amps']} A, {f['status']})"
            for f in feeders[:12]
        )
        blocks.append(f"Feeder inventory for {name}: {listing}.")

    if outages:
        p1 = [o for o in outages if o["severity"] == "P1"]
        restored = [o for o in outages if o["mttr_minutes"] is not None]
        avg_mttr = sum(o["mttr_minutes"] for o in restored) / len(restored) if restored else 0
        causes = {}
        for o in outages:
            causes[o["cause"]] = causes.get(o["cause"], 0) + 1
        top_cause = max(causes, key=causes.get)
        blocks.append(
            f"Over the trailing twelve months, {name} recorded {len(outages)} outage events "
            f"across its feeders, including {len(p1)} classified P1. Average restoration time "
            f"(MTTR) across restored events was {avg_mttr:.0f} minutes. The most frequent "
            f"reported cause was {top_cause} ({causes[top_cause]} events)."
        )
    else:
        blocks.append(
            f"No outage events were recorded against {name} feeders during the trailing "
            f"twelve-month reporting window."
        )

    blocks.append(
        f"Operational notes: switching at {name} must be coordinated with the control center "
        f"under the standard switching-order procedure. "
        + (
            "As this is a transmission-class facility, all switching orders require verification "
            "by a second qualified person before execution."
            if sub["substation_type"] == "transmission"
            else "Distribution-class switching may be executed by a single qualified operator "
                 "once the switching order is approved."
        )
        + (
            f" This substation is currently in {sub['status']} state; confirm clearance status "
            f"before dispatching any crew to site."
            if sub["status"] != "in_service" else ""
        )
    )

    return f"substation-profile-{sub['substation_id']:03d}", f"Substation Operating Profile — {name}", blocks


def build_outage_postmortem(rng: random.Random, out: dict, dispatches: list[dict]) -> tuple[str, str, list[str]]:
    oid = out["outage_id"]
    started = out["started_at"]
    restored = out["restored_at"]
    duration = f"{out['mttr_minutes']} minutes" if out["mttr_minutes"] is not None else "not yet restored"

    blocks = [
        f"Outage #{oid} was recorded on feeder {out['feeder_name']} (served from substation "
        f"{out['substation_name']}, {out['region']} region) beginning "
        f"{started.strftime('%Y-%m-%d at %H:%M UTC')}. The event was classified severity "
        f"{out['severity']} and affected approximately {out['customers_affected']} customers. "
        f"Reported cause: {out['cause']}.",
    ]

    if restored is not None:
        blocks.append(
            f"Service was restored at {restored.strftime('%Y-%m-%d %H:%M UTC')}, giving a total "
            f"restoration time of {duration}. "
            + (
                f"This is within the {out['severity']} target restoration window."
                if out["mttr_minutes"] is not None
                and out["mttr_minutes"] <= {"P1": 120, "P2": 240, "P3": 480, "P4": 1440}[out["severity"]]
                else f"This exceeded the {out['severity']} target restoration window and was "
                     f"flagged for reliability review."
            )
        )
    else:
        blocks.append(
            f"As of this report, outage #{oid} remains open with no restoration timestamp "
            f"recorded. Restoration status should be confirmed with the control center."
        )

    if dispatches:
        d = dispatches[0]
        blocks.append(
            f"Crew response: {d['crew_name']} was dispatched at "
            f"{d['dispatched_at'].strftime('%Y-%m-%d %H:%M UTC')} with a recorded response time of "
            f"{d['response_time_mins']} minutes to on-site arrival."
            + (
                " Response time met the P1 target of 45 minutes for normal weather conditions."
                if out["severity"] == "P1" and d["response_time_mins"] <= 45
                else " Response time is logged for trending against severity targets."
            )
        )

    blocks.append(f"Root cause summary: {out['rca_summary']}")

    blocks.append(
        "Corrective actions: "
        + rng.choice([
            "vegetation management cycle for this feeder segment has been advanced to the "
            "next available crew window.",
            "the affected line section is scheduled for a follow-up patrol to confirm no "
            "residual damage remains.",
            "fault indicator placement along this feeder is being reviewed to reduce future "
            "patrol time.",
            "no corrective action beyond standard repair was identified; the event is "
            "considered closed.",
            "this feeder has been added to the worst-performing-feeder watch list for the "
            "current reporting year.",
        ])
    )

    return f"outage-postmortem-{oid:04d}", f"Outage Post-Mortem Report — Outage #{oid}", blocks


def build_transformer_inspection(rng: random.Random, tr: dict, alarms: list[dict]) -> tuple[str, str, list[str]]:
    tid = tr["transformer_id"]
    age = datetime.now().year - tr["install_year"]

    blocks = [
        f"Asset inspection record for transformer {tr['asset_tag']}, a "
        f"{float(tr['capacity_mva']):.2f} MVA unit manufactured by {tr['manufacturer']} and "
        f"installed in {tr['install_year']} ({age} years in service). The unit is connected to "
        f"feeder {tr['feeder_name']} at substation {tr['substation_name']} and its current "
        f"status is {tr['status']}.",

        f"Visual inspection findings: oil level and condition — {rng.choice(OIL_CONDITIONS)}. "
        f"Bushings — {rng.choice(BUSHING_CONDITIONS)}. Cooling fans and radiators were "
        f"{rng.choice(['operating normally', 'operating normally with minor debris accumulation', 'cycling correctly under load'])}. "
        f"Grounding connections were verified intact. Inspection performed by a "
        f"{rng.choice(INSPECTOR_TITLES)} under weather conditions: {rng.choice(WEATHER)}.",
    ]

    if alarms:
        by_type: dict[str, int] = {}
        for a in alarms:
            by_type[a["alarm_type"]] = by_type.get(a["alarm_type"], 0) + 1
        summary = ", ".join(f"{k} ({v})" for k, v in sorted(by_type.items()))
        unresolved = [a for a in alarms if not a["resolved"]]
        blocks.append(
            f"SCADA alarm history for {tr['asset_tag']} over the trailing twelve months shows "
            f"{len(alarms)} recorded alarms: {summary}. "
            f"{len(unresolved)} of these remain unresolved."
        )
        thermal = sum(v for k, v in by_type.items()
                      if k in ("transformer overtemperature", "overcurrent"))
        if thermal >= 3:
            blocks.append(
                f"Note: this unit has recorded {thermal} overtemperature/overcurrent alarms, "
                f"meeting the threshold of three or more thermal alarms in a rolling 90-day "
                f"window that triggers a capacity upgrade review. {tr['asset_tag']} should be "
                f"referred to planning for a loading study."
            )
        else:
            blocks.append(
                f"Thermal alarm count ({thermal}) is below the capacity-upgrade-review threshold. "
                f"No loading study is required at this time."
            )
    else:
        blocks.append(
            f"No SCADA alarms were recorded against {tr['asset_tag']} during the trailing "
            f"twelve-month window. The unit is considered to be in nominal service."
        )

    blocks.append(
        "Recommended follow-up: "
        + (
            f"unit is {tr['status']}; confirm clearance and tagout status before any further work."
            if tr["status"] != "in_service"
            else rng.choice([
                "no action required; return to standard inspection rotation.",
                "schedule dissolved gas analysis at the next routine maintenance window.",
                "re-inspect at the standard interval; no defects requiring intervention were found.",
                "clean bushing surfaces during the next scheduled outage window.",
            ])
        )
    )

    return f"transformer-inspection-{tid:04d}", f"Transformer Inspection Report — {tr['asset_tag']}", blocks


def build_feeder_summary(rng: random.Random, fd: dict, outages: list[dict],
                         meters: dict) -> tuple[str, str, list[str]]:
    fid = fd["feeder_id"]
    total_meters = sum(meters.values())

    blocks = [
        f"Reliability summary for feeder {fd['name']}, a {fd['voltage_kv']} kV distribution "
        f"circuit rated {fd['capacity_amps']} A, served from substation {fd['substation_name']} "
        f"in the {fd['region']} region. Current feeder status: {fd['status']}.",

        f"Customer base: this feeder serves {total_meters} metered service points — "
        + ", ".join(f"{v} {k}" for k, v in sorted(meters.items()) if v)
        + ". Only active and inactive meters count toward customers-affected figures in "
          "reliability reporting; disconnected services are excluded.",
    ]

    if outages:
        sustained = [o for o in outages if o["mttr_minutes"] is not None]
        total_cust_minutes = sum(
            (o["customers_affected"] or 0) * (o["mttr_minutes"] or 0) for o in sustained
        )
        by_sev: dict[str, int] = {}
        for o in outages:
            by_sev[o["severity"]] = by_sev.get(o["severity"], 0) + 1
        sev_line = ", ".join(f"{k}: {v}" for k, v in sorted(by_sev.items()))
        avg_mttr = (
            sum(o["mttr_minutes"] for o in sustained) / len(sustained) if sustained else 0
        )
        blocks.append(
            f"Outage history over the trailing twelve months: {len(outages)} recorded events "
            f"({sev_line}). Average restoration time was {avg_mttr:.0f} minutes across "
            f"{len(sustained)} restored events, for a total of approximately "
            f"{total_cust_minutes:,} customer-minutes of interruption on this feeder."
        )
        causes: dict[str, int] = {}
        for o in outages:
            causes[o["cause"]] = causes.get(o["cause"], 0) + 1
        ranked = sorted(causes.items(), key=lambda kv: -kv[1])
        blocks.append(
            "Cause breakdown: "
            + "; ".join(f"{c} — {n} event(s)" for c, n in ranked)
            + ". "
            + (
                f"With {ranked[0][1]} events attributed to {ranked[0][0]}, this feeder is a "
                f"candidate for targeted reliability work addressing that failure mode."
                if ranked[0][1] >= 3
                else "No single failure mode dominates this feeder's outage history."
            )
        )
    else:
        blocks.append(
            f"No sustained outage events were recorded on {fd['name']} during the trailing "
            f"twelve-month window. This feeder is performing above the system average and "
            f"requires no targeted reliability intervention."
        )

    return f"feeder-summary-{fid:04d}", f"Feeder Reliability Summary — {fd['name']}", blocks


def build_regional_report(rng: random.Random, region: str, month: str, year: int,
                          outages: list[dict], customers: int) -> tuple[str, str, list[str]]:
    sustained = [o for o in outages if o["mttr_minutes"] is not None]
    cust_minutes = sum((o["customers_affected"] or 0) * (o["mttr_minutes"] or 0) for o in sustained)
    cust_interrupted = sum(o["customers_affected"] or 0 for o in sustained)
    saidi = cust_minutes / customers if customers else 0
    saifi = cust_interrupted / customers if customers else 0
    caidi = saidi / saifi if saifi else 0
    by_sev: dict[str, int] = {}
    for o in outages:
        by_sev[o["severity"]] = by_sev.get(o["severity"], 0) + 1

    blocks = [
        f"Monthly reliability report for the {region} region covering {month} {year}. "
        f"This report summarizes distribution reliability performance for all feeders in the "
        f"region and is prepared for internal review and regulatory submission.",

        f"During the reporting period the region recorded {len(outages)} outage events "
        + (", ".join(f"{v} at severity {k}" for k, v in sorted(by_sev.items())) if by_sev else "none")
        + f". {len(sustained)} events were restored within the period, accounting for "
          f"approximately {cust_minutes:,} customer-minutes of interruption across "
          f"{cust_interrupted:,} customer interruptions.",

        f"Calculated reliability indices for {region} in {month} {year}: "
        f"SAIDI {saidi:.2f} minutes per customer, SAIFI {saifi:.3f} interruptions per customer, "
        f"CAIDI {caidi:.1f} minutes per interrupted customer. These figures are calculated "
        f"against a served base of {customers:,} metered customers in the region.",

        "Methodology note: SAIDI is the sum of customers affected multiplied by outage duration "
        "in minutes, divided by total customers served. SAIFI is total customer interruptions "
        "divided by total customers served. CAIDI is SAIDI divided by SAIFI. Momentary "
        "interruptions cleared by automatic reclose are excluded from these figures and are "
        "tracked separately under MAIFI. Major event days, where daily SAIDI exceeds the "
        "statistical threshold, are excluded from routine reporting and reported separately.",
    ]

    return (
        f"regional-report-{region.lower().replace(' ', '-')}-{year}-{MONTHS.index(month) + 1:02d}",
        f"Monthly Reliability Report — {region} — {month} {year}",
        blocks,
    )


def build_storm_report(rng: random.Random, idx: int, name: str, region: str,
                       outages: list[dict]) -> tuple[str, str, list[str]]:
    total_cust = sum(o["customers_affected"] or 0 for o in outages)
    restored = [o for o in outages if o["mttr_minutes"] is not None]
    peak_mttr = max((o["mttr_minutes"] for o in restored), default=0)
    crews = rng.randint(8, 40)
    mutual_aid = rng.randint(0, 25)

    blocks = [
        f"After-action report for {name}, a declared major weather event affecting the "
        f"{region} region. The Emergency Response Plan was activated ahead of the event based "
        f"on forecast severity, with the control center staffed to 24-hour coverage and "
        f"{crews} internal crews placed on standby."
        + (
            f" An additional {mutual_aid} mutual-aid crews were requested from neighboring "
            f"utilities under existing mutual-assistance agreements."
            if mutual_aid else
            " Internal crew capacity was sufficient; no mutual-aid request was issued."
        ),

        f"Event impact: {len(outages)} outage events were attributed to this storm across the "
        f"region, affecting approximately {total_cust:,} customers in total. The longest single "
        f"restoration took {peak_mttr} minutes. During the event, dispatch operated under the "
        f"storm model — prioritizing greatest customers restored per crew-hour rather than "
        f"strict severity order — with safety hazards made safe and cordoned in the interim.",

        "Estimated restoration times were communicated to the public at regional granularity "
        "rather than per-customer, consistent with the Emergency Response Plan, since "
        "individual ERTs cannot be reliably estimated until damage assessment is complete. "
        "Outage tickets in areas with a common cause were consolidated into single restoration "
        "jobs to avoid duplicate crew dispatch.",

        "Debrief findings: "
        + rng.choice([
            "staging site pre-positioning materially reduced average crew travel time; the "
            "same staging plan should be reused for comparable forecasts.",
            "damage assessment was delayed by limited early patrol capacity; adding dedicated "
            "assessment crews ahead of restoration crews is recommended for future events.",
            "mutual-aid crews required longer than expected onboarding on local switching "
            "procedures; a pre-event briefing packet is recommended.",
            "communication loss to several substation RTUs extended the time required to "
            "confirm feeder status; RTU backup power should be reviewed.",
        ])
        + " This event was classified a major event day for reliability reporting purposes and "
          "its outages are excluded from routine blue-sky SAIDI and SAIFI figures.",
    ]

    return f"storm-after-action-{idx:02d}", f"Storm After-Action Report — {name}", blocks


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main() -> None:
    ap = argparse.ArgumentParser(description="Generate GridOps operational document corpus")
    ap.add_argument("--force", action="store_true", help="Wipe the output directory first")
    args = ap.parse_args()

    rng = random.Random(SEED)

    if OUT_DIR.exists() and args.force:
        shutil.rmtree(OUT_DIR)
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    conn = psycopg2.connect(settings.database_url)
    written: list[str] = []

    substations = fetch_all(conn, "SELECT * FROM substations ORDER BY substation_id")
    feeders = fetch_all(
        conn,
        """SELECT f.*, s.name AS substation_name, s.region
           FROM feeders f JOIN substations s ON s.substation_id = f.substation_id
           ORDER BY f.feeder_id""",
    )
    transformers = fetch_all(
        conn,
        """SELECT t.*, f.name AS feeder_name, s.name AS substation_name
           FROM transformers t
           JOIN feeders f ON f.feeder_id = t.feeder_id
           JOIN substations s ON s.substation_id = f.substation_id
           ORDER BY t.transformer_id""",
    )
    outages = fetch_all(
        conn,
        """SELECT o.*, f.name AS feeder_name, s.name AS substation_name, s.region
           FROM outages o
           JOIN feeders f ON f.feeder_id = o.feeder_id
           JOIN substations s ON s.substation_id = f.substation_id
           ORDER BY o.outage_id""",
    )
    alarms = fetch_all(conn, "SELECT * FROM scada_alarms ORDER BY alarm_id")
    dispatches = fetch_all(conn, "SELECT * FROM crew_dispatch_logs ORDER BY log_id")
    meter_rows = fetch_all(
        conn, "SELECT feeder_id, customer_type, status, count(*) AS n FROM meters GROUP BY 1,2,3"
    )
    conn.close()

    # --- index by parent ---
    feeders_by_sub: dict[int, list[dict]] = {}
    for f in feeders:
        feeders_by_sub.setdefault(f["substation_id"], []).append(f)
    tx_by_sub: dict[int, list[dict]] = {}
    tx_by_feeder: dict[int, list[dict]] = {}
    for t in transformers:
        tx_by_feeder.setdefault(t["feeder_id"], []).append(t)
    for f in feeders:
        for t in tx_by_feeder.get(f["feeder_id"], []):
            tx_by_sub.setdefault(f["substation_id"], []).append(t)
    outages_by_feeder: dict[int, list[dict]] = {}
    for o in outages:
        outages_by_feeder.setdefault(o["feeder_id"], []).append(o)
    outages_by_sub: dict[int, list[dict]] = {}
    for f in feeders:
        for o in outages_by_feeder.get(f["feeder_id"], []):
            outages_by_sub.setdefault(f["substation_id"], []).append(o)
    alarms_by_tx: dict[int, list[dict]] = {}
    for a in alarms:
        alarms_by_tx.setdefault(a["source_transformer_id"], []).append(a)
    dispatch_by_outage: dict[int, list[dict]] = {}
    for d in dispatches:
        dispatch_by_outage.setdefault(d["outage_id"], []).append(d)
    meters_by_feeder: dict[int, dict[str, int]] = {}
    for m in meter_rows:
        meters_by_feeder.setdefault(m["feeder_id"], {})
        key = m["customer_type"]
        meters_by_feeder[m["feeder_id"]][key] = (
            meters_by_feeder[m["feeder_id"]].get(key, 0) + m["n"]
        )

    # --- 1. substation profiles ---
    for sub in substations[:N_SUBSTATION_PROFILES]:
        slug, title, blocks = build_substation_profile(
            rng, sub,
            feeders_by_sub.get(sub["substation_id"], []),
            tx_by_sub.get(sub["substation_id"], []),
            outages_by_sub.get(sub["substation_id"], []),
        )
        written.append(emit(rng, slug, title, blocks))

    # --- 2. outage post-mortems (prioritise P1/P2, then fill) ---
    ranked_outages = sorted(outages, key=lambda o: ({"P1": 0, "P2": 1, "P3": 2, "P4": 3}[o["severity"]], o["outage_id"]))
    for out in ranked_outages[:N_OUTAGE_POSTMORTEMS]:
        slug, title, blocks = build_outage_postmortem(
            rng, out, dispatch_by_outage.get(out["outage_id"], [])
        )
        written.append(emit(rng, slug, title, blocks))

    # --- 3. transformer inspections (prefer units with alarm history) ---
    ranked_tx = sorted(transformers, key=lambda t: -len(alarms_by_tx.get(t["transformer_id"], [])))
    for tr in ranked_tx[:N_TRANSFORMER_INSPECTIONS]:
        slug, title, blocks = build_transformer_inspection(
            rng, tr, alarms_by_tx.get(tr["transformer_id"], [])
        )
        written.append(emit(rng, slug, title, blocks))

    # --- 4. feeder reliability summaries (prefer feeders with outage history) ---
    ranked_fd = sorted(feeders, key=lambda f: -len(outages_by_feeder.get(f["feeder_id"], [])))
    for fd in ranked_fd[:N_FEEDER_SUMMARIES]:
        slug, title, blocks = build_feeder_summary(
            rng, fd,
            outages_by_feeder.get(fd["feeder_id"], []),
            meters_by_feeder.get(fd["feeder_id"], {}),
        )
        written.append(emit(rng, slug, title, blocks))

    # --- 5. regional monthly reports ---
    regions = sorted({s["region"] for s in substations})
    outages_by_region_month: dict[tuple[str, int, int], list[dict]] = {}
    for o in outages:
        key = (o["region"], o["started_at"].year, o["started_at"].month)
        outages_by_region_month.setdefault(key, []).append(o)
    customers_by_region: dict[str, int] = {}
    for f in feeders:
        customers_by_region[f["region"]] = customers_by_region.get(f["region"], 0) + sum(
            meters_by_feeder.get(f["feeder_id"], {}).values()
        )
    report_keys = sorted(outages_by_region_month.keys())[:N_REGIONAL_REPORTS]
    for region, year, month_num in report_keys:
        slug, title, blocks = build_regional_report(
            rng, region, MONTHS[month_num - 1], year,
            outages_by_region_month[(region, year, month_num)],
            customers_by_region.get(region, 1),
        )
        written.append(emit(rng, slug, title, blocks))

    # --- 6. storm after-action reports ---
    storm_outages = [o for o in outages if o["cause"] == "storm damage"]
    for i, name in enumerate(STORM_NAMES[:N_STORM_REPORTS], start=1):
        region = regions[i % len(regions)]
        subset = [o for o in storm_outages if o["region"] == region] or storm_outages[:5]
        slug, title, blocks = build_storm_report(rng, i, name, region, subset)
        written.append(emit(rng, slug, title, blocks))

    # --- report ---
    by_ext: dict[str, int] = {}
    for n in written:
        by_ext[Path(n).suffix] = by_ext.get(Path(n).suffix, 0) + 1
    print(f"Wrote {len(written)} documents to {OUT_DIR}")
    for ext, n in sorted(by_ext.items()):
        print(f"  {ext:6} {n:>4}")


if __name__ == "__main__":
    main()
