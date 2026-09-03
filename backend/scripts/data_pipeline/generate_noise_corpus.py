"""Generate the out-of-domain distractor corpus (`seed/docs/noisy_data/`).

The original design for this project called for a ~95% noise / 5% signal
knowledge base so that each advanced retrieval technique has to *earn its
place*: with a corpus that is all signal, naive top-k retrieval looks good and
reranking/CRAG have nothing to rescue.

Two flavours of distractor are generated, deliberately:

* **near-domain** - telecom / water utility / datacenter / HVAC / rail / fleet
  operations. These share the operational vocabulary of grid ops ("outage",
  "alarm", "maintenance window", "crew", "capacity", "P1") without answering
  any grid question. This is the noise that actually defeats naive dense
  retrieval and forces the cross-encoder to discriminate.
* **far-domain** - recipes, travel, HR policy, gardening, personal finance.
  Trivially separable by embedding, included so the corpus is not uniformly
  adversarial.

Usage:
    uv run python scripts/data_pipeline/generate_noise_corpus.py [--count N] [--force]

Writes: seed/docs/noisy_data/
"""

from __future__ import annotations

import argparse
import random
import shutil
from pathlib import Path

OUT_DIR = Path(__file__).parent.parent.parent / "seed" / "docs" / "noisy_data"
SEED = 20260903
DEFAULT_COUNT = 1200

# Near-domain noise gets the larger share - it is the noise that matters.
NEAR_SHARE = 0.7

FORMAT_WEIGHTS = {".txt": 55, ".md": 22, ".html": 18, ".docx": 4, ".pdf": 1}

SEVERITIES = ["P1", "P2", "P3", "P4"]
CREWS = ["Team Alpha", "Team Bravo", "Team Delta", "Night Shift", "Regional Crew 3"]

NEAR_DOMAINS = {
    "telecom": {
        "asset": ["cell site", "fiber span", "core router", "microwave link", "DSLAM cabinet"],
        "tag": "SITE",
        "fault": [
            "fiber cut",
            "power module failure",
            "backhaul congestion",
            "antenna misalignment",
            "cooling fan failure",
        ],
        "metric": ["packet loss", "jitter", "backhaul utilization", "call setup success rate"],
    },
    "water": {
        "asset": ["lift station", "treatment basin", "distribution main", "booster pump", "reservoir"],
        "tag": "WTR",
        "fault": [
            "main break",
            "pump seal failure",
            "turbidity excursion",
            "chlorine residual low",
            "SCADA telemetry loss",
        ],
        "metric": ["flow rate", "residual chlorine", "turbidity", "reservoir level"],
    },
    "datacenter": {
        "asset": ["rack row", "CRAC unit", "UPS bank", "core switch", "chiller loop"],
        "tag": "DC",
        "fault": [
            "PSU failure",
            "thermal excursion",
            "UPS transfer fault",
            "uplink flap",
            "humidity out of band",
        ],
        "metric": ["inlet temperature", "PUE", "rack load", "uplink utilization"],
    },
    "hvac": {
        "asset": ["air handler", "chiller", "VAV box", "cooling tower", "boiler"],
        "tag": "AHU",
        "fault": [
            "belt failure",
            "filter differential high",
            "valve actuator stuck",
            "refrigerant low",
            "controller offline",
        ],
        "metric": ["supply air temperature", "static pressure", "chilled water delta-T"],
    },
    "rail": {
        "asset": ["signal block", "interlocking", "level crossing", "point machine", "track circuit"],
        "tag": "SIG",
        "fault": [
            "point machine failure",
            "track circuit occupancy fault",
            "crossing gate fault",
            "axle counter reset",
            "cable theft",
        ],
        "metric": ["headway", "dwell time", "signal availability"],
    },
    "fleet": {
        "asset": ["service van", "bucket truck", "utility trailer", "pool vehicle"],
        "tag": "VEH",
        "fault": [
            "brake wear",
            "hydraulic leak",
            "battery failure",
            "telematics unit offline",
            "tyre replacement",
        ],
        "metric": ["odometer", "fuel efficiency", "engine hours"],
    },
}

FAR_TOPICS = {
    "recipe": (
        "Recipe: {dish}",
        [
            "This {dish} recipe serves {n} and takes about {t} minutes from start to finish. "
            "It is a weeknight staple because most of the work is hands-off.",
            "Ingredients: {ing}. Substitutions are forgiving; the ratio of fat to acid matters "
            "more than the specific brand.",
            "Method: preheat, combine the dry ingredients, then fold in the wet ingredients "
            "without overworking the mixture. Rest for {t2} minutes before serving.",
            "Leftovers keep for {d} days refrigerated. Reheat gently - aggressive heat will "
            "split the sauce and dull the aromatics.",
        ],
    ),
    "travel": (
        "Travel Notes: {place}",
        [
            "{place} is best visited outside peak season, when accommodation costs drop by "
            "roughly {n} percent and the main sights are walkable without queuing.",
            "Getting around is straightforward: regional trains run every {t} minutes and a "
            "day pass costs less than two single fares.",
            "Food is the highlight. Look for family-run places away from the central square, "
            "where the menu is short and changes with what the market had that morning.",
            "Budget roughly {d} days to see the area properly without rushing between sites.",
        ],
    ),
    "hr_policy": (
        "Employee Handbook: {topic} Policy",
        [
            "This policy sets out the {topic} entitlement for all permanent employees and "
            "takes effect from the start of the next review cycle.",
            "Employees accrue {n} days per calendar year, pro-rated for part-time staff. "
            "Requests should be submitted at least {t} working days in advance.",
            "Approval sits with the line manager. Where a request is declined, the reason "
            "must be recorded and the employee offered an alternative window.",
            "Unused entitlement may be carried over up to a maximum of {d} days; anything "
            "beyond that lapses at year end.",
        ],
    ),
    "gardening": (
        "Growing Guide: {plant}",
        [
            "{plant} prefers well-drained soil and at least {n} hours of direct sun. In "
            "heavier clay soils, raise the bed or work in coarse grit before planting.",
            "Sow seed at a depth of roughly {t} millimetres and thin seedlings once the "
            "first true leaves appear.",
            "Water deeply but infrequently. Shallow daily watering encourages surface roots "
            "and makes the plant less resilient in a dry spell.",
            "Expect first harvest around {d} weeks after transplanting, earlier under cover.",
        ],
    ),
    "finance": (
        "Personal Finance Note: {topic}",
        [
            "Understanding {topic} matters because small differences compound. A {n} percent "
            "difference in annual cost is not noticeable in one year and substantial over twenty.",
            "The mechanics are simple: contributions are made pre-tax, growth is sheltered, "
            "and tax is paid on withdrawal at the prevailing rate.",
            "A common mistake is holding cash inside a tax-sheltered account for {t} years, "
            "which wastes the shelter on an asset that generates little taxable growth.",
            "Review allocations roughly every {d} months - more often than that tends to "
            "encourage trading rather than discipline.",
        ],
    ),
}

DISHES = [
    "braised short rib", "lemon risotto", "chickpea stew", "miso aubergine",
    "buttermilk cornbread", "green curry", "mushroom ragu", "shakshuka",
]
PLACES = [
    "Porto", "Ljubljana", "Kanazawa", "Tbilisi", "Valparaiso", "Gdansk",
    "Hobart", "Bergen", "Oaxaca", "Trieste",
]
HR_TOPICS = [
    "Annual Leave", "Remote Working", "Parental Leave", "Expense Claims",
    "Study Leave", "Sick Pay",
]
PLANTS = ["garlic", "runner beans", "rhubarb", "chard", "leeks", "raspberries", "tomatillo"]
FIN_TOPICS = [
    "index fund fees", "emergency funds", "tax-sheltered accounts",
    "debt prioritisation", "rebalancing",
]
INGREDIENTS = [
    "olive oil, onion, garlic, tinned tomatoes, smoked paprika",
    "butter, shallot, arborio rice, white wine, parmesan",
    "coconut milk, lemongrass, ginger, lime, coriander",
    "flour, buttermilk, cornmeal, baking soda, honey",
]


def write_txt(path: Path, title: str, blocks: list[str]) -> None:
    path.write_text(title + "\n\n" + "\n\n".join(blocks), encoding="utf-8")


def write_md(path: Path, title: str, blocks: list[str]) -> None:
    path.write_text(f"# {title}\n\n" + "\n\n".join(blocks), encoding="utf-8")


def write_html(path: Path, title: str, blocks: list[str]) -> None:
    paras = "\n".join(f"  <p>{b}</p>" for b in blocks)
    path.write_text(
        "<!doctype html>\n<html><head><meta charset='utf-8'>"
        f"<title>{title}</title></head>\n<body>\n  <h1>{title}</h1>\n{paras}\n</body></html>\n",
        encoding="utf-8",
    )


def write_docx(path: Path, title: str, blocks: list[str]) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    for b in blocks:
        doc.add_paragraph(b)
    doc.save(str(path))


def write_pdf(path: Path, title: str, blocks: list[str]) -> None:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(path), pagesize=LETTER)
    flow = [Paragraph(title, styles["Title"]), Spacer(1, 12)]
    for b in blocks:
        flow.append(Paragraph(b, styles["BodyText"]))
        flow.append(Spacer(1, 8))
    doc.build(flow)


WRITERS = {
    ".txt": write_txt,
    ".md": write_md,
    ".html": write_html,
    ".docx": write_docx,
    ".pdf": write_pdf,
}


def build_near(rng: random.Random, domain: str, idx: int) -> tuple[str, str, list[str]]:
    d = NEAR_DOMAINS[domain]
    asset = rng.choice(d["asset"])
    tag = f"{d['tag']}-{rng.randint(1000, 9999)}"
    fault = rng.choice(d["fault"])
    metric = rng.choice(d["metric"])
    sev = rng.choice(SEVERITIES)
    crew = rng.choice(CREWS)
    mins = rng.randint(12, 460)
    affected = rng.randint(5, 9000)

    title = f"{domain.title()} Operations Incident {tag}"
    blocks = [
        f"Incident {tag} was raised against {asset} at {rng.randint(0, 23):02d}:"
        f"{rng.randint(0, 59):02d} and classified severity {sev}. The reported fault was "
        f"{fault}. Approximately {affected} subscribers were impacted for the duration of "
        f"the event.",
        f"Monitoring showed {metric} drifting outside its normal band roughly "
        f"{rng.randint(5, 90)} minutes before the alarm cleared threshold, which suggests "
        f"the condition developed gradually rather than as a step change.",
        f"{crew} was dispatched and restored service after {mins} minutes. Root cause was "
        f"confirmed as {fault}; the failed component was replaced from van stock and the "
        f"spare re-ordered against the standing supply agreement.",
        f"Follow-up: the maintenance window for this {asset} has been brought forward, and "
        f"the {metric} alarm threshold is under review because the current setting did not "
        f"give useful early warning. No wider network impact was identified.",
    ]
    return f"{domain}-incident-{idx:05d}", title, blocks


def build_far(rng: random.Random, topic: str, idx: int) -> tuple[str, str, list[str]]:
    title_tpl, para_tpls = FAR_TOPICS[topic]
    subs = {
        "dish": rng.choice(DISHES),
        "place": rng.choice(PLACES),
        "topic": rng.choice(HR_TOPICS if topic == "hr_policy" else FIN_TOPICS),
        "plant": rng.choice(PLANTS),
        "ing": rng.choice(INGREDIENTS),
        "n": rng.randint(2, 40),
        "t": rng.randint(5, 90),
        "t2": rng.randint(5, 30),
        "d": rng.randint(2, 30),
    }
    title = title_tpl.format(**subs)
    blocks = [p.format(**subs) for p in para_tpls]
    return f"{topic}-{idx:05d}", title, blocks


def main() -> None:
    ap = argparse.ArgumentParser(description="Generate out-of-domain distractor corpus")
    ap.add_argument("--count", type=int, default=DEFAULT_COUNT)
    ap.add_argument("--force", action="store_true", help="Wipe noisy_data/ first")
    args = ap.parse_args()

    rng = random.Random(SEED)

    if OUT_DIR.exists() and args.force:
        for p in OUT_DIR.iterdir():
            if p.name == ".gitkeep":
                continue
            if p.is_file():
                p.unlink()
            else:
                shutil.rmtree(p)
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    n_near = int(args.count * NEAR_SHARE)
    n_far = args.count - n_near
    near_domains = list(NEAR_DOMAINS)
    far_topics = list(FAR_TOPICS)

    counts: dict[str, int] = {}
    for i in range(n_near):
        domain = near_domains[i % len(near_domains)]
        slug, title, blocks = build_near(rng, domain, i)
        ext = rng.choices(list(FORMAT_WEIGHTS), weights=list(FORMAT_WEIGHTS.values()))[0]
        WRITERS[ext](OUT_DIR / f"{slug}{ext}", title, blocks)
        counts[ext] = counts.get(ext, 0) + 1

    for i in range(n_far):
        topic = far_topics[i % len(far_topics)]
        slug, title, blocks = build_far(rng, topic, i)
        ext = rng.choices(list(FORMAT_WEIGHTS), weights=list(FORMAT_WEIGHTS.values()))[0]
        WRITERS[ext](OUT_DIR / f"{slug}{ext}", title, blocks)
        counts[ext] = counts.get(ext, 0) + 1

    print(f"Wrote {n_near} near-domain + {n_far} far-domain = {args.count} noise docs")
    for ext, n in sorted(counts.items()):
        print(f"  {ext:6} {n:>5}")


if __name__ == "__main__":
    main()
